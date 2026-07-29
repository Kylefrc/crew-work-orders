"""Low-level python-docx helpers for crew work orders.

Brand-neutral building blocks: page setup, body text, checkboxes, safety
banner, photo grid, and the raw XML cell utilities the branded layer
(brand.py) is built on. Nothing in this file knows about any company.
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAFETY_RED = RGBColor(0xC0, 0x50, 0x4D)

FONT = 'Arial'


# ─── Cell / paragraph XML utilities ─────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set cell background fill (hex string, no #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='AAAAAA'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tc_pr.append(borders)


def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tc_pr.append(mar)


def set_col_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(width_twips))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)


def add_para_bottom_border(para, color='888888', sz='6'):
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ─── Document-level helpers ─────────────────────────────────────────────────

def create_doc():
    """New Document with narrow margins that fit a full day table per page."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    return doc


def add_body(doc, text, bold=False, italic=False, size=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT
    run.bold = bold
    run.italic = italic
    return para


def add_checkbox(doc, text):
    """Checkbox item (Unicode ballot box) crews tick with a pen."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run('☐  ' + text)
    run.font.size = Pt(10)
    run.font.name = FONT
    return para


def add_safety_banner(doc, flags):
    """Red banner listing active safety flags. Silent when there are none."""
    if not flags:
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run('⚠  SAFETY FLAGS: ' + ' | '.join(flags))
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT
    run.font.color.rgb = SAFETY_RED


def add_photo_grid(doc, photo_paths, captions=None, cols=2):
    """Captioned photo grid. Missing files render an inline note, never crash."""
    if not photo_paths:
        return
    rows_needed = (len(photo_paths) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed * 2, cols=cols)
    for idx, path in enumerate(photo_paths):
        r = (idx // cols) * 2
        c = idx % cols
        cell = table.rows[r].cells[c]
        p = cell.paragraphs[0]
        if os.path.exists(path):
            p.add_run().add_picture(path, width=Inches(3.0))
        else:
            run = p.add_run(f'[Photo not found: {os.path.basename(path)}]')
            run.font.size = Pt(8)
            run.italic = True
        if captions and idx < len(captions) and captions[idx]:
            cap = table.rows[r + 1].cells[c].paragraphs[0]
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(captions[idx])
            run.font.size = Pt(8)
            run.font.name = FONT
            run.italic = True
