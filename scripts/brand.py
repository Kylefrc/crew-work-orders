#!/usr/bin/env python3
"""Config-driven branding for crew work orders.

Reads a company config (simple `- key: value` markdown, see config/company.md)
and exposes the branded document builders: banner header, section headings,
key-value info blocks, tables, checklists, footer. All colors derive from two
hex values in the config, so adopting this package for a new company is a
config edit — no code changes.

    from brand import Brand
    b = Brand.load('config/company.md')
    b.header(doc, 'HVAC WORK ORDER', 'Alder Residence — Portland, OR')
"""

import os

from docx.shared import Pt, Inches, RGBColor

from docx_helpers import (
    add_checkbox, add_para_bottom_border, cell_margins,
    set_cell_bg, set_cell_borders, set_col_width, FONT,
)

WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DEFAULTS = {
    'company_name': 'Your Company Name',
    'license_line': '',
    'prepared_by': '',
    'crew_contact': '',
    'logo': '',
    'color_primary': '1F4E79',   # deep slate blue — swap for your brand
    'color_accent': '6E8FAF',
    'vent_ratio': '300',
    'qa_label': 'QA',
    'checklist_dirs': '../checklists',
}


def _blend(hex_color, other, t):
    """Blend two hex colors; t=0 returns hex_color, t=1 returns other."""
    a = [int(hex_color[i:i + 2], 16) for i in (0, 2, 4)]
    b = [int(other[i:i + 2], 16) for i in (0, 2, 4)]
    return ''.join(f'{round(x + (y - x) * t):02X}' for x, y in zip(a, b))


def _rgb(hex_color):
    return RGBColor(*(int(hex_color[i:i + 2], 16) for i in (0, 2, 4)))


def load_config(path):
    """Parse `- key: value` (or `key: value`) lines; ignore blanks and # comments."""
    cfg = dict(DEFAULTS)
    with open(path, encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            if line.startswith('- '):
                line = line[2:]
            if ':' not in line:
                continue
            key, value = line.split(':', 1)
            key = key.strip().lower().replace(' ', '_')
            if key in DEFAULTS:
                cfg[key] = value.strip()
    cfg['_dir'] = os.path.dirname(os.path.abspath(path))
    return cfg


class Brand:
    def __init__(self, cfg):
        self.cfg = cfg
        self.name = cfg['company_name']
        self.primary_hex = cfg['color_primary'].lstrip('#').upper()
        self.accent_hex = cfg['color_accent'].lstrip('#').upper()
        self.primary = _rgb(self.primary_hex)
        # Derived palette: tinted rows and borders that always harmonize
        self.pale_hex = _blend(self.accent_hex, 'FFFFFF', 0.82)
        self.border_hex = _blend(self.accent_hex, 'FFFFFF', 0.45)
        self.alt_row_hex = 'F7F7F2'
        logo = cfg.get('logo', '').strip()
        self.logo = os.path.join(cfg['_dir'], logo) if logo else ''

    @classmethod
    def load(cls, config_path):
        return cls(load_config(config_path))

    def checklist_dirs(self):
        """Checklist library directories, resolved relative to the config file.
        Later directories override earlier ones by checklist key."""
        dirs = []
        for entry in self.cfg['checklist_dirs'].split(','):
            entry = entry.strip()
            if entry:
                dirs.append(os.path.normpath(os.path.join(self.cfg['_dir'], entry)))
        return dirs

    # ── Branded builders ────────────────────────────────────────────────────

    def header(self, doc, doc_type, subtitle=''):
        """Full-width primary-color banner: logo (or wordmark) left, doc type right."""
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0]
        c0, c1 = row.cells[0], row.cells[1]
        for c in (c0, c1):
            set_cell_bg(c, self.primary_hex)
            set_cell_borders(c, self.primary_hex)
            cell_margins(c, top=140, bottom=140, left=160, right=160)
        set_col_width(c0, 3200)
        set_col_width(c1, 6160)
        p0 = c0.paragraphs[0]
        if self.logo and os.path.exists(self.logo):
            p0.add_run().add_picture(self.logo, width=Inches(1.5))
        else:
            r = p0.add_run(self.name.upper())
            r.bold = True
            r.font.size = Pt(14)
            r.font.name = FONT
            r.font.color.rgb = WHITE
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(doc_type)
        r1.bold = True
        r1.font.size = Pt(18)
        r1.font.name = FONT
        r1.font.color.rgb = WHITE
        if subtitle:
            p1b = c1.add_paragraph()
            r2 = p1b.add_run(subtitle)
            r2.font.size = Pt(10)
            r2.font.name = FONT
            r2.font.color.rgb = WHITE
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def heading(self, doc, text):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = FONT
        run.font.color.rgb = self.primary
        add_para_bottom_border(para, self.accent_hex, '8')
        return para

    def two_col(self, doc, rows):
        table = doc.add_table(rows=len(rows), cols=2)
        for i, (label, value) in enumerate(rows):
            r = table.rows[i]
            c0, c1 = r.cells[0], r.cells[1]
            fill = self.pale_hex if i % 2 == 0 else 'FFFFFF'
            for c in (c0, c1):
                set_cell_bg(c, fill)
                set_cell_borders(c, self.border_hex)
                cell_margins(c)
            set_col_width(c0, 2600)
            set_col_width(c1, 6760)
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(label)
            r0.bold = True
            r0.font.size = Pt(10)
            r0.font.name = FONT
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(str(value))
            r1.font.size = Pt(10)
            r1.font.name = FONT
        return table

    def table(self, doc, headers, rows, widths=None):
        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        for j, hdr in enumerate(headers):
            cell = table.rows[0].cells[j]
            set_cell_bg(cell, self.primary_hex)
            set_cell_borders(cell, self.primary_hex)
            cell_margins(cell)
            if widths:
                set_col_width(cell, widths[j])
            p = cell.paragraphs[0]
            r = p.add_run(hdr)
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = FONT
            r.font.color.rgb = WHITE
        for i, row_data in enumerate(rows):
            fill = self.alt_row_hex if i % 2 == 0 else 'FFFFFF'
            for j, val in enumerate(row_data):
                cell = table.rows[i + 1].cells[j]
                set_cell_bg(cell, fill)
                set_cell_borders(cell, self.border_hex)
                cell_margins(cell)
                if widths:
                    set_col_width(cell, widths[j])
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(9)
                r.font.name = FONT
        return table

    def checklist(self, doc, title, items):
        self.heading(doc, title)
        for item in items:
            add_checkbox(doc, item)

    def footer(self, doc):
        """Divider + company line. Work is credited to a person, never software."""
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        add_para_bottom_border(para, self.accent_hex, '6')
        parts = [self.name]
        if self.cfg['license_line']:
            parts.append(self.cfg['license_line'])
        if self.cfg['prepared_by']:
            parts.append(f"Prepared by: {self.cfg['prepared_by']}")
        p2 = doc.add_paragraph()
        run = p2.add_run('  ·  '.join(parts))
        run.font.size = Pt(9)
        run.font.name = FONT
        run.font.color.rgb = self.primary
